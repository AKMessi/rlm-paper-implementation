"""
Document Processor for RLM
==========================

Handles loading, parsing, and chunking of very large documents.
Supports multiple formats: PDF, TXT, DOCX, Markdown, etc.
"""

import io
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Union, Callable
from dataclasses import dataclass
import logging

import aiofiles
from pypdf import PdfReader

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """Represents a processed document."""
    content: Union[str, List[str], Dict[str, Any]]  # Can be string, chunks, or structured
    metadata: Dict[str, Any]
    source: str
    doc_type: str
    total_chars: int
    num_chunks: int = 1
    
    def get_context_for_rlm(self) -> Any:
        """Get the content in a format suitable for RLM processing."""
        return self.content


class DocumentProcessor:
    """
    Processor for large documents.
    
    Handles:
    - File loading and parsing
    - Text extraction from various formats
    - Intelligent chunking strategies
    - Metadata extraction
    """
    
    def __init__(
        self,
        chunk_size: int = 100000,  # ~100KB per chunk (paper mentions ~500K chars for sub-LM)
        chunk_overlap: int = 1000,
        respect_boundaries: bool = True,  # Try to respect paragraph/section boundaries
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.respect_boundaries = respect_boundaries
        
    async def load_document(
        self, 
        file_path: Union[str, Path],
        content: Optional[bytes] = None,
        filename: Optional[str] = None
    ) -> Document:
        """
        Load and process a document.
        
        Args:
            file_path: Path to the file (or identifier if content provided)
            content: Optional raw bytes content
            filename: Original filename (for extension detection)
            
        Returns:
            Processed Document object
        """
        file_path = Path(file_path)
        filename = filename or file_path.name
        extension = Path(filename).suffix.lower()
        
        logger.info(f"Loading document: {filename}")
        
        # Read content if not provided
        if content is None:
            async with aiofiles.open(file_path, 'rb') as f:
                content = await f.read()
        
        # Parse based on file type
        if extension == '.pdf':
            doc = await self._parse_pdf(content, filename)
        elif extension in ['.docx', '.doc']:
            doc = await self._parse_docx(content, filename)
        elif extension in ['.txt', '.md', '.markdown']:
            doc = await self._parse_text(content, filename)
        elif extension in ['.py', '.js', '.java', '.cpp', '.c', '.h', '.ts', '.tsx', '.jsx']:
            doc = await self._parse_code(content, filename)
        elif extension in ['.json']:
            doc = await self._parse_json(content, filename)
        else:
            # Try as text
            try:
                doc = await self._parse_text(content, filename)
            except:
                logger.warning(f"Unknown file type: {extension}, treating as binary text")
                doc = await self._parse_binary_as_text(content, filename)
        
        return doc
    
    async def _parse_pdf(self, content: bytes, filename: str) -> Document:
        """Parse PDF document."""
        logger.info(f"Parsing PDF: {filename}")
        
        try:
            pdf_file = io.BytesIO(content)
            reader = PdfReader(pdf_file)
            
            pages = []
            total_text = ""
            
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    pages.append({
                        "page_num": i + 1,
                        "text": text,
                        "char_count": len(text)
                    })
                    total_text += text + "\n\n"
            
            # Chunk intelligently
            chunks = self._chunk_text_intelligently(total_text)
            
            metadata = {
                "filename": filename,
                "num_pages": len(reader.pages),
                "file_type": "pdf",
            }
            
            return Document(
                content=chunks,  # Return as list of chunks for RLM
                metadata=metadata,
                source=filename,
                doc_type="pdf",
                total_chars=len(total_text),
                num_chunks=len(chunks)
            )
            
        except Exception as e:
            logger.error(f"Error parsing PDF {filename}: {str(e)}")
            raise
    
    async def _parse_docx(self, content: bytes, filename: str) -> Document:
        """Parse DOCX document."""
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available, falling back to text extraction")
            return await self._parse_binary_as_text(content, filename)
        
        logger.info(f"Parsing DOCX: {filename}")
        
        try:
            doc_file = io.BytesIO(content)
            doc = DocxDocument(doc_file)
            
            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Also extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    tables_text.append(" | ".join(row_text))
            
            all_text = "\n\n".join(paragraphs + tables_text)
            chunks = self._chunk_text_intelligently(all_text)
            
            metadata = {
                "filename": filename,
                "num_paragraphs": len(paragraphs),
                "num_tables": len(doc.tables),
                "file_type": "docx",
            }
            
            return Document(
                content=chunks,
                metadata=metadata,
                source=filename,
                doc_type="docx",
                total_chars=len(all_text),
                num_chunks=len(chunks)
            )
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {filename}: {str(e)}")
            return await self._parse_binary_as_text(content, filename)
    
    async def _parse_text(self, content: bytes, filename: str) -> Document:
        """Parse plain text or markdown."""
        logger.info(f"Parsing text: {filename}")
        
        # Detect encoding
        text = self._decode_bytes(content)
        
        # For markdown, try to respect header boundaries
        is_markdown = filename.endswith(('.md', '.markdown'))
        
        if is_markdown and self.respect_boundaries:
            chunks = self._chunk_markdown(text)
        else:
            chunks = self._chunk_text_intelligently(text)
        
        metadata = {
            "filename": filename,
            "file_type": "markdown" if is_markdown else "text",
        }
        
        return Document(
            content=chunks,
            metadata=metadata,
            source=filename,
            doc_type="text",
            total_chars=len(text),
            num_chunks=len(chunks)
        )
    
    async def _parse_code(self, content: bytes, filename: str) -> Document:
        """Parse code files with structure awareness."""
        logger.info(f"Parsing code: {filename}")
        
        text = self._decode_bytes(content)
        
        # For code, try to respect function/class boundaries
        chunks = self._chunk_code(text, filename)
        
        metadata = {
            "filename": filename,
            "file_type": "code",
            "language": Path(filename).suffix[1:],
        }
        
        return Document(
            content=chunks,
            metadata=metadata,
            source=filename,
            doc_type="code",
            total_chars=len(text),
            num_chunks=len(chunks)
        )
    
    async def _parse_json(self, content: bytes, filename: str) -> Document:
        """Parse JSON files."""
        import json
        
        logger.info(f"Parsing JSON: {filename}")
        
        text = self._decode_bytes(content)
        
        try:
            data = json.loads(text)
            # Pretty print for easier processing
            formatted = json.dumps(data, indent=2)
            
            # If it's a list, chunk by items
            if isinstance(data, list):
                chunks = []
                current_chunk = []
                current_size = 0
                
                for item in data:
                    item_str = json.dumps(item, indent=2)
                    if current_size + len(item_str) > self.chunk_size and current_chunk:
                        chunks.append("[\n" + ",\n".join(current_chunk) + "\n]")
                        current_chunk = []
                        current_size = 0
                    current_chunk.append(item_str)
                    current_size += len(item_str)
                
                if current_chunk:
                    chunks.append("[\n" + ",\n".join(current_chunk) + "\n]")
            else:
                chunks = self._chunk_text_intelligently(formatted)
            
            metadata = {
                "filename": filename,
                "file_type": "json",
            }
            
            return Document(
                content=chunks,
                metadata=metadata,
                source=filename,
                doc_type="json",
                total_chars=len(formatted),
                num_chunks=len(chunks)
            )
            
        except json.JSONDecodeError:
            # Treat as text
            return await self._parse_text(content, filename)
    
    async def _parse_binary_as_text(self, content: bytes, filename: str) -> Document:
        """Fallback: parse binary content as text."""
        text = self._decode_bytes(content, errors='replace')
        chunks = self._chunk_text_intelligently(text)
        
        return Document(
            content=chunks,
            metadata={"filename": filename, "file_type": "unknown"},
            source=filename,
            doc_type="unknown",
            total_chars=len(text),
            num_chunks=len(chunks)
        )
    
    def _decode_bytes(self, content: bytes, errors: str = 'strict') -> str:
        """Decode bytes to string, trying multiple encodings."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        # Last resort
        return content.decode('utf-8', errors=errors)
    
    def _chunk_text_intelligently(self, text: str) -> List[str]:
        """
        Chunk text intelligently, respecting boundaries when possible.
        
        Strategy:
        1. Try to split on paragraph boundaries (\n\n)
        2. Try to split on line boundaries (\n)
        3. Fall back to character boundaries
        """
        if len(text) <= self.chunk_size:
            return [text]
        
        chunks = []
        current_chunk = ""
        current_size = 0
        
        # Split on paragraphs first
        paragraphs = text.split('\n\n')
        
        for para in paragraphs:
            para_size = len(para) + 2  # +2 for \n\n
            
            if current_size + para_size > self.chunk_size:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    
                    # Add overlap from previous chunk
                    if self.chunk_overlap > 0:
                        overlap_text = current_chunk[-self.chunk_overlap:]
                        current_chunk = overlap_text + "\n\n" + para
                        current_size = len(current_chunk)
                    else:
                        current_chunk = para
                        current_size = len(para)
                else:
                    # Paragraph itself is larger than chunk_size
                    # Split on lines
                    chunks.extend(self._chunk_by_lines(para))
                    current_chunk = ""
                    current_size = 0
            else:
                if current_chunk:
                    current_chunk += "\n\n" + para
                else:
                    current_chunk = para
                current_size += para_size
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _chunk_by_lines(self, text: str) -> List[str]:
        """Chunk text by lines when paragraphs are too long."""
        lines = text.split('\n')
        chunks = []
        current_chunk = ""
        
        for line in lines:
            if len(current_chunk) + len(line) + 1 > self.chunk_size:
                if current_chunk:
                    chunks.append(current_chunk)
                    # Overlap
                    if self.chunk_overlap > 0:
                        current_chunk = current_chunk[-self.chunk_overlap:] + "\n" + line
                    else:
                        current_chunk = line
                else:
                    # Line itself is too long, force split
                    for i in range(0, len(line), self.chunk_size):
                        chunks.append(line[i:i + self.chunk_size])
                    current_chunk = ""
            else:
                current_chunk += "\n" + line if current_chunk else line
        
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    def _chunk_markdown(self, text: str) -> List[str]:
        """Chunk markdown respecting header boundaries."""
        # Split on headers (### ## #)
        header_pattern = r'^(#{1,6}\s+.+)$'
        
        sections = []
        current_section = ""
        
        for line in text.split('\n'):
            if re.match(header_pattern, line):
                if current_section:
                    sections.append(current_section)
                current_section = line + "\n"
            else:
                current_section += line + "\n"
        
        if current_section:
            sections.append(current_section)
        
        # Now chunk sections, respecting boundaries
        chunks = []
        current_chunk = ""
        
        for section in sections:
            if len(current_chunk) + len(section) > self.chunk_size:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = section
            else:
                current_chunk += "\n" + section
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks if chunks else [text]
    
    def _chunk_code(self, text: str, filename: str) -> List[str]:
        """Chunk code respecting function/class boundaries."""
        # Simple pattern matching for common languages
        # This could be enhanced with AST parsing
        
        # Pattern for Python/JavaScript/Java/C++ style functions/classes
        block_patterns = [
            r'(def\s+\w+\s*\([^)]*\):\s*(?:\n\s*(?:"""|\'\'\')[^\n]*(?:\n[^"\']*(?:"""|\'\'\'))?\s*\n)?(?:\n\s+[^\n]*)*)',  # Python
            r'(function\s+\w+\s*\([^)]*\)\s*\{[^}]*\})',  # JavaScript function
            r'(class\s+\w+[^}]*\})',  # JavaScript/Java/C++ class
            r'(public|private|protected)?\s*[\w<>\[\]]+\s+\w+\s*\([^)]*\)\s*\{[^}]*\}',  # Java/C++ method
        ]
        
        # For simplicity, use line-based chunking with awareness of indentation
        lines = text.split('\n')
        chunks = []
        current_chunk = []
        current_size = 0
        
        i = 0
        while i < len(lines):
            line = lines[i]
            line_size = len(line) + 1
            
            # Check if this is the start of a new block (based on indentation)
            if line.strip() and not line.startswith(' ') and not line.startswith('\t'):
                # Potential block start
                if current_chunk and current_size + line_size > self.chunk_size:
                    chunks.append('\n'.join(current_chunk))
                    current_chunk = []
                    current_size = 0
            
            current_chunk.append(line)
            current_size += line_size
            
            if current_size > self.chunk_size and len(current_chunk) > 1:
                # Save all but last line
                chunks.append('\n'.join(current_chunk[:-1]))
                current_chunk = [current_chunk[-1]]
                current_size = len(current_chunk[0])
            
            i += 1
        
        if current_chunk:
            chunks.append('\n'.join(current_chunk))
        
        return chunks if chunks else [text]
    
    def combine_documents(self, documents: List[Document]) -> Document:
        """
        Combine multiple documents into a single document for RLM processing.
        
        This is useful when the user uploads multiple files and wants to query
        across all of them.
        """
        all_chunks = []
        total_chars = 0
        sources = []
        
        for doc in documents:
            if isinstance(doc.content, list):
                # Add source prefix to each chunk
                for i, chunk in enumerate(doc.content):
                    prefixed = f"=== {doc.source} (chunk {i+1}/{len(doc.content)}) ===\n\n{chunk}"
                    all_chunks.append(prefixed)
            else:
                prefixed = f"=== {doc.source} ===\n\n{str(doc.content)}"
                all_chunks.append(prefixed)
            
            total_chars += doc.total_chars
            sources.append(doc.source)
        
        metadata = {
            "sources": sources,
            "num_documents": len(documents),
            "combined": True,
        }
        
        return Document(
            content=all_chunks,
            metadata=metadata,
            source="combined_documents",
            doc_type="combined",
            total_chars=total_chars,
            num_chunks=len(all_chunks)
        )
